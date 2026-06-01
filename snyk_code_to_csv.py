#!/usr/bin/env python3
"""
snyk_code_to_csv.py

Run `snyk code test` and export the findings to a CSV and/or a PDF report that
includes the full source-to-sink data flow (with real code snippets) and the
remediation advice ("how to fix") for every issue.

Output is built from Snyk Code's SARIF output, which is the only format that
carries both the data-flow steps (codeFlows/threadFlows) and the per-rule
remediation guidance (rule.help.markdown).

Usage:
    # Scan a project and write snyk-results.csv
    ./snyk_code_to_csv.py /path/to/project

    # Produce a PDF report instead (or both)
    ./snyk_code_to_csv.py /path/to/project --format pdf
    ./snyk_code_to_csv.py /path/to/project --format both -o report

    # Re-use an existing SARIF file instead of scanning
    ./snyk_code_to_csv.py --sarif-input results.sarif --project-root /path/to/project

    # Custom output path and pass extra args through to the Snyk CLI
    ./snyk_code_to_csv.py . -o findings.csv -- --org=my-org --severity-threshold=medium

CSV export uses only the standard library. PDF export additionally requires
`reportlab` (pip3 install reportlab); it is imported lazily so CSV works without it.

Snyk CLI exit codes: 0 = no issues, 1 = issues found (both are success here),
2 = scan error.
"""

import argparse
import csv
import datetime
import json
import os
import re
import subprocess
import sys
import tempfile

# SARIF `level` -> Snyk severity label (Snyk Code uses error/warning/note)
LEVEL_TO_SEVERITY = {"error": "High", "warning": "Medium", "note": "Low"}
# SCA adds a Critical tier; sort Critical first
SEVERITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}

CSV_COLUMNS = [
    "scan_type",
    "severity",
    "issue_title",
    "rule_id",
    "cwe",
    "cve",
    "cvss",
    "package",
    "fixed_in",
    "exploit_maturity",
    "priority_score",
    "file",
    "line",
    "message",
    "source",
    "sink",
    "data_flow",
    "remediation",
    "autofixable",
    "fingerprint",
]


def blank_row():
    """A row dict with every column present (so writers can index any key)."""
    return {col: "" for col in CSV_COLUMNS}


def validate_input_path(raw_path):
    """Resolve a user-supplied SARIF path and confine it to the working directory.

    The path is canonicalised (resolving symlinks and ``..``) and then required
    to live inside the current working directory tree. This blocks path
    traversal -- the input cannot escape the project to read arbitrary files
    such as ``../../etc/passwd`` -- and rejects anything that is not a real
    regular file.
    """
    base = os.path.realpath(os.getcwd())
    resolved = os.path.realpath(raw_path)
    try:
        contained = os.path.commonpath([base, resolved]) == base
    except ValueError:  # different drives (Windows) -> not contained
        contained = False
    if not contained:
        raise SystemExit(
            f"SARIF input must be inside the working directory ({base}): {raw_path}"
        )
    if not os.path.isfile(resolved):
        raise SystemExit(f"SARIF input is not a readable file: {raw_path}")
    return resolved


def run_snyk_code(target, extra_args):
    """Run `snyk code test` writing SARIF to a temp file. Returns the SARIF path."""
    tmp = tempfile.NamedTemporaryFile(
        prefix="snyk-code-", suffix=".sarif", delete=False
    )
    tmp.close()
    cmd = ["snyk", "code", "test", target, f"--sarif-file-output={tmp.name}"]
    cmd.extend(extra_args)
    print(f"Running: {' '.join(cmd)}", file=sys.stderr)
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    # 0 = clean, 1 = issues found -> both fine. Anything else and no sarif = error.
    if result.returncode not in (0, 1) and not os.path.getsize(tmp.name):
        sys.stderr.write(result.stdout.decode("utf-8", "replace"))
        raise SystemExit(f"snyk code test failed (exit {result.returncode})")
    if not os.path.getsize(tmp.name):
        sys.stderr.write(result.stdout.decode("utf-8", "replace"))
        raise SystemExit("snyk produced no SARIF output (no results or scan error)")
    return tmp.name


def run_snyk_sca(target, extra_args):
    """Run `snyk test` (open source) writing JSON to a temp file.

    Returns the JSON path, or None if there were no supported manifests to scan
    (so `--type both` can still report Code results).
    """
    tmp = tempfile.NamedTemporaryFile(prefix="snyk-sca-", suffix=".json", delete=False)
    tmp.close()
    cmd = ["snyk", "test", target, "--all-projects", f"--json-file-output={tmp.name}"]
    cmd.extend(extra_args)
    print(f"Running: {' '.join(cmd)}", file=sys.stderr)
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
    # 0 = clean, 1 = vulns found -> both fine. 3 = no supported projects found.
    if not os.path.getsize(tmp.name):
        output = result.stdout.decode("utf-8", "replace")
        if result.returncode in (0, 3):
            print("No open-source dependencies found to scan (SCA skipped).", file=sys.stderr)
            return None
        sys.stderr.write(output)
        raise SystemExit(f"snyk test failed (exit {result.returncode})")
    return tmp.name


class SourceReader:
    """Reads and caches source files so we can attach real code to flow steps."""

    def __init__(self, project_root):
        self.root = project_root
        self.cache = {}

    def _load(self, uri):
        if uri not in self.cache:
            path = os.path.join(self.root, uri)
            try:
                with open(path, "r", encoding="utf-8", errors="replace") as fh:
                    self.cache[uri] = fh.read().splitlines()
            except OSError:
                self.cache[uri] = None
        return self.cache[uri]

    def line(self, uri, line_no):
        lines = self._load(uri)
        if not lines or line_no < 1 or line_no > len(lines):
            return ""
        return lines[line_no - 1].strip()


def build_rule_index(run):
    """Map ruleId/index -> rule object for title, CWE and remediation lookup."""
    rules = run.get("tool", {}).get("driver", {}).get("rules", [])
    by_id = {r.get("id"): r for r in rules}
    return rules, by_id


def get_rule(result, rules, by_id):
    rid = result.get("ruleId")
    if rid in by_id:
        return by_id[rid]
    idx = result.get("ruleIndex")
    if isinstance(idx, int) and 0 <= idx < len(rules):
        return rules[idx]
    return {}


def physical(loc):
    """Return (uri, start_line) for a SARIF location, or (None, None)."""
    phys = loc.get("physicalLocation", {})
    uri = phys.get("artifactLocation", {}).get("uri")
    line = phys.get("region", {}).get("startLine")
    return uri, line


def extract_flow_steps(result):
    """Flatten the primary code flow into ordered, de-duplicated steps."""
    flows = result.get("codeFlows") or []
    if not flows:
        return []
    thread_flows = flows[0].get("threadFlows") or []
    if not thread_flows:
        return []
    steps = []
    prev = None
    for tf_loc in thread_flows[0].get("locations", []):
        uri, line = physical(tf_loc.get("location", {}))
        if uri is None or line is None:
            continue
        key = (uri, line)
        if key == prev:  # collapse consecutive repeats of the same line
            continue
        prev = key
        steps.append((uri, line))
    return steps


def render_flow(steps, reader):
    """Render the full source-to-sink path with code snippets, one step per line."""
    lines = []
    multi_file = len({uri for uri, _ in steps}) > 1
    for i, (uri, line_no) in enumerate(steps, 1):
        code = reader.line(uri, line_no)
        ref = f"{uri}:{line_no}" if multi_file else f"L{line_no}"
        lines.append(f"{i}. {ref}  |  {code}")
    return "\n".join(lines)


def endpoint(steps, reader, which):
    """Format the source (first) or sink (last) step as 'file:line | code'."""
    if not steps:
        return ""
    uri, line_no = steps[0] if which == "source" else steps[-1]
    code = reader.line(uri, line_no)
    return f"{uri}:{line_no}  |  {code}".rstrip(" |")


def result_to_row(result, rules, by_id, reader):
    rule = get_rule(result, rules, by_id)
    props = result.get("properties", {}) or {}

    uri, line = physical(result.get("locations", [{}])[0]) if result.get("locations") else (None, None)
    severity = LEVEL_TO_SEVERITY.get(result.get("level"), result.get("level", ""))
    cwes = ", ".join(rule.get("properties", {}).get("cwe", []) or [])

    steps = extract_flow_steps(result)

    row = blank_row()
    row.update({
        "scan_type": "Code",
        "severity": severity,
        "issue_title": rule.get("shortDescription", {}).get("text", rule.get("name", "")),
        "rule_id": result.get("ruleId", ""),
        "cwe": cwes,
        "priority_score": props.get("priorityScore", ""),
        "file": uri or "",
        "line": line or "",
        "message": result.get("message", {}).get("text", ""),
        "source": endpoint(steps, reader, "source"),
        "sink": endpoint(steps, reader, "sink"),
        "data_flow": render_flow(steps, reader),
        "remediation": (rule.get("help", {}).get("markdown")
                        or rule.get("help", {}).get("text", "")).strip(),
        "autofixable": props.get("isAutofixable", ""),
        "fingerprint": (result.get("fingerprints", {}) or {}).get("identity", "")
        or (result.get("fingerprints", {}) or {}).get("0", ""),
    })
    return row


def sarif_to_rows(sarif, reader):
    rows = []
    for run in sarif.get("runs", []):
        rules, by_id = build_rule_index(run)
        for result in run.get("results", []):
            rows.append(result_to_row(result, rules, by_id, reader))
    return rows


# --- SCA (open source) parsing ---------------------------------------------

def _upgrade_advice(vuln):
    """A one-line fix sentence from the SCA upgrade/patch/fixed-in data."""
    pkg = vuln.get("packageName", "the package")
    fixed_in = vuln.get("fixedIn") or []
    if vuln.get("isUpgradable"):
        # upgradePath[0] is often `false` (the top dep can't be pinned); the
        # first string entry is the direct dependency to bump.
        direct = next((p for p in vuln.get("upgradePath", []) if isinstance(p, str)), None)
        if direct:
            target = f" (fixed in `{pkg}` {', '.join(fixed_in)})" if fixed_in else ""
            return f"Upgrade `{direct}`{target}."
    if fixed_in:
        return f"Upgrade `{pkg}` to {fixed_in[0]} or later."
    if vuln.get("isPatchable"):
        return f"A Snyk patch is available for `{pkg}`."
    return f"No fixed version is available yet for `{pkg}`."


def vuln_to_row(vuln, manifest):
    pkg = vuln.get("packageName", "")
    version = vuln.get("version", "")
    chain = vuln.get("from", []) or []
    ids = vuln.get("identifiers", {}) or {}
    cve = ", ".join(ids.get("CVE", []) or [])
    cwe = ", ".join(ids.get("CWE", []) or [])

    flow = "\n".join(f"{i}. {'  ' * min(i - 1, 8)}{dep}" for i, dep in enumerate(chain, 1))
    fix = _upgrade_advice(vuln)
    intro = " → ".join(chain)
    remediation = f"**Fix:** {fix}\n\n**Introduced through:** {intro}\n\n{vuln.get('description', '')}".strip()

    severity = (vuln.get("severityWithCritical") or vuln.get("severity") or "").capitalize()

    row = blank_row()
    row.update({
        "scan_type": "SCA",
        "severity": severity,
        "issue_title": vuln.get("title", ""),
        "rule_id": vuln.get("id", ""),
        "cwe": cwe,
        "cve": cve,
        "cvss": vuln.get("cvssScore", ""),
        "package": f"{pkg}@{version}" if version else pkg,
        "fixed_in": ", ".join(vuln.get("fixedIn") or []),
        "exploit_maturity": vuln.get("exploit", ""),
        "file": manifest,
        "message": f"{vuln.get('title', '')} in {pkg}@{version}".strip(),
        "source": chain[0] if chain else "",
        "sink": chain[-1] if chain else f"{pkg}@{version}",
        "data_flow": flow,
        "remediation": remediation,
        "fingerprint": vuln.get("id", ""),
    })
    return row


def sca_json_to_rows(sca):
    """Parse `snyk test --json` output (single object or --all-projects array)."""
    projects = sca if isinstance(sca, list) else [sca]
    rows = []
    for proj in projects:
        if not isinstance(proj, dict) or proj.get("error"):
            continue
        manifest = proj.get("displayTargetFile") or proj.get("targetFile") or proj.get("projectName", "")
        seen = set()
        for vuln in proj.get("vulnerabilities", []) or []:
            vid = vuln.get("id")
            # the same vuln id can appear once per dependency path; keep the first
            if vid in seen:
                continue
            seen.add(vid)
            rows.append(vuln_to_row(vuln, manifest))
    return rows


def sort_rows(rows):
    rows.sort(key=lambda r: (SEVERITY_ORDER.get(r["severity"], 9),
                             r["scan_type"], r["file"],
                             r["line"] if isinstance(r["line"], int) else 0))
    return rows


def severity_counts(rows):
    counts = {}
    for r in rows:
        counts[r["severity"]] = counts.get(r["severity"], 0) + 1
    return counts


def summary_text(rows):
    counts = severity_counts(rows)
    return ", ".join(f"{counts[s]} {s}"
                     for s in ("Critical", "High", "Medium", "Low") if s in counts)


def write_csv(rows, path, _meta=None):
    # _meta is accepted for a uniform writer signature; CSV doesn't use it.
    with open(path, "w", encoding="utf-8", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=CSV_COLUMNS)
        writer.writeheader()
        writer.writerows(rows)


# --- PDF export ------------------------------------------------------------

SEVERITY_COLORS = {"Critical": "#9b1c1c", "High": "#d1352c",
                   "Medium": "#d98c00", "Low": "#3b7dd8"}


def finding_meta(r):
    """Per-finding metadata table rows and flow-section label, by scan type."""
    if r["scan_type"] == "SCA":
        info = [
            ("Package", r["package"]),
            ("Manifest", r["file"]),
            ("CVE", r["cve"] or "-"),
            ("CVSS", str(r["cvss"]) or "-"),
            ("CWE", r["cwe"] or "-"),
            ("Fixed in", r["fixed_in"] or "-"),
            ("Exploit maturity", r["exploit_maturity"] or "-"),
            ("Snyk ID", r["rule_id"]),
        ]
        return info, "Dependency path (introduced through)"
    info = [
        ("File", f'{r["file"]}:{r["line"]}'),
        ("Rule", r["rule_id"]),
        ("CWE", r["cwe"] or "-"),
        ("Priority score", str(r["priority_score"]) or "-"),
        ("Auto-fixable", str(r["autofixable"]) or "-"),
    ]
    return info, "Data flow (source → sink)"


def _escape(text):
    return (str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _inline_md(text):
    """Convert inline markdown to reportlab's mini-markup (after escaping)."""
    out = _escape(text)
    out = re.sub(r"`([^`]+)`",
                 r'<font face="Courier" backColor="#f0f0f0">\1</font>', out)
    out = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", out)
    out = re.sub(r"\[([^\]]+)\]\(([^)]+)\)",
                 r'<a href="\2" color="#3b7dd8">\1</a>', out)
    return out


def _wrap_mono(text, width=108):
    """Hard-wrap monospace text so long lines (e.g. SQL) don't overflow the page."""
    import textwrap
    wrapped = []
    for line in text.split("\n"):
        if len(line) <= width:
            wrapped.append(line)
        else:
            wrapped.extend(textwrap.wrap(
                line, width=width, subsequent_indent="      ",
                break_long_words=True, break_on_hyphens=False) or [""])
    return "\n".join(wrapped)


def _markdown_to_flowables(md, styles):
    from reportlab.platypus import Paragraph, Spacer, Preformatted, ListFlowable, ListItem

    flow, bullets, para = [], [], []

    def flush_para():
        if para:
            flow.append(Paragraph(_inline_md(" ".join(s.strip() for s in para)), styles["body"]))
            para.clear()

    def flush_bullets():
        if bullets:
            items = [ListItem(Paragraph(_inline_md(b), styles["body"])) for b in bullets]
            flow.append(ListFlowable(items, bulletType="bullet", leftIndent=14))
            flow.append(Spacer(1, 4))
            bullets.clear()

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.lstrip().startswith("```"):
            flush_para(); flush_bullets()
            code, i = [], i + 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                code.append(lines[i]); i += 1
            flow.append(Preformatted(_wrap_mono("\n".join(code)), styles["code"]))
            flow.append(Spacer(1, 4))
            i += 1
            continue
        stripped = line.strip()
        if not stripped:
            flush_para(); flush_bullets()
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if h:
            flush_para(); flush_bullets()
            style = styles["h2"] if len(h.group(1)) <= 2 else styles["h3"]
            flow.append(Paragraph(_inline_md(h.group(2)), style))
            i += 1
            continue
        b = re.match(r"^[-*]\s+(.*)", stripped)
        if b:
            flush_para()
            bullets.append(b.group(1))
            i += 1
            continue
        flush_bullets()
        para.append(line)
        i += 1
    flush_para(); flush_bullets()
    return flow


def _pdf_styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=base["Title"], fontSize=20, spaceAfter=4),
        "meta": ParagraphStyle("meta", parent=base["Normal"], fontSize=9,
                               textColor=colors.HexColor("#666666")),
        "finding": ParagraphStyle("finding", parent=base["Heading2"], fontSize=13, spaceBefore=8, spaceAfter=2),
        "h2": ParagraphStyle("h2", parent=base["Heading3"], fontSize=11, spaceBefore=6, spaceAfter=2),
        "h3": ParagraphStyle("h3", parent=base["Heading4"], fontSize=10, spaceBefore=4, spaceAfter=2),
        "label": ParagraphStyle("label", parent=base["Normal"], fontSize=9, fontName="Helvetica-Bold",
                                textColor=colors.HexColor("#444444"), spaceBefore=4),
        "body": ParagraphStyle("body", parent=base["Normal"], fontSize=9.5, leading=13),
        "code": ParagraphStyle("code", parent=base["Code"], fontSize=7.2, leading=9,
                               backColor=colors.HexColor("#f5f5f5"), borderPadding=4,
                               textColor=colors.HexColor("#1a1a1a")),
    }


def write_pdf(rows, path, meta):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                         Preformatted, Table, TableStyle, HRFlowable)
    except ImportError:
        raise SystemExit(
            "PDF export requires the 'reportlab' package. Install it with:\n"
            "    pip3 install reportlab"
        )

    styles = _pdf_styles()
    doc = SimpleDocTemplate(path, pagesize=A4, title="Snyk Code Report",
                            leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm)
    story = []

    # Header
    story.append(Paragraph("Snyk Security Report", styles["title"]))
    story.append(Paragraph(
        f"Project: {_escape(meta['project'])}<br/>"
        f"Scan: {_escape(meta['scan'])}<br/>"
        f"Generated: {meta['date']}<br/>"
        f"Total issues: {len(rows)} ({summary_text(rows) or 'none'})", styles["meta"]))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", color=colors.HexColor("#cccccc")))

    for idx, r in enumerate(rows, 1):
        sev = r["severity"]
        info, flow_label = finding_meta(r)
        story.append(Spacer(1, 8))
        story.append(Paragraph(
            f'<font color="{SEVERITY_COLORS.get(sev, "#666666")}">[{sev}]</font> '
            f'<font color="#888888">[{r["scan_type"]}]</font> '
            f'{idx}. {_escape(r["issue_title"])}', styles["finding"]))

        tbl = Table([[k, str(v)] for k, v in info], colWidths=[30 * mm, 144 * mm])
        tbl.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#444444")),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#eeeeee")),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(tbl)

        if r["message"]:
            story.append(Paragraph(_inline_md(r["message"]), styles["body"]))

        story.append(Paragraph(_escape(flow_label), styles["label"]))
        story.append(Preformatted(_wrap_mono(r["data_flow"] or "(none)"), styles["code"]))

        if r["remediation"]:
            story.append(Paragraph("Remediation", styles["label"]))
            story.extend(_markdown_to_flowables(r["remediation"], styles))

        story.append(Spacer(1, 6))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#dddddd")))

    if not rows:
        story.append(Spacer(1, 12))
        story.append(Paragraph("No issues found.", styles["body"]))

    doc.build(story)


# --- DOCX export -----------------------------------------------------------

_MD_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _docx_add_inline(paragraph, text):
    """Add a markdown line to a docx paragraph, honouring **bold**, `code`, links."""
    from docx.shared import RGBColor

    for part in _MD_INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if link:
                run = paragraph.add_run(f"{link.group(1)} ({link.group(2)})")
                run.font.color.rgb = RGBColor(0x3B, 0x7D, 0xD8)
            else:
                paragraph.add_run(part)


def _docx_add_markdown(doc, md):
    """Render the remediation markdown into a docx document."""
    from docx.shared import Pt

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.lstrip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run(lines[i])
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                i += 1
            i += 1
            continue
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if h:
            doc.add_heading(h.group(2), level=min(len(h.group(1)) + 1, 4))
            i += 1
            continue
        b = re.match(r"^[-*]\s+(.*)", stripped)
        if b:
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline(p, b.group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        _docx_add_inline(p, stripped)
        i += 1


def write_docx(rows, path, meta):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise SystemExit(
            "DOCX export requires the 'python-docx' package. Install it with:\n"
            "    pip3 install python-docx"
        )

    doc = Document()
    doc.add_heading("Snyk Security Report", level=0)
    intro = doc.add_paragraph()
    intro.add_run("Project: ").bold = True
    intro.add_run(meta["project"])
    intro.add_run("\nScan: ").bold = True
    intro.add_run(meta["scan"])
    intro.add_run("\nGenerated: ").bold = True
    intro.add_run(meta["date"])
    intro.add_run("\nTotal issues: ").bold = True
    intro.add_run(f"{len(rows)} ({summary_text(rows) or 'none'})")

    for idx, r in enumerate(rows, 1):
        sev = r["severity"]
        info, flow_label = finding_meta(r)
        heading = doc.add_heading(level=1)
        hexcol = SEVERITY_COLORS.get(sev, "#666666")
        sev_run = heading.add_run(f"[{sev}] ")
        sev_run.font.color.rgb = RGBColor.from_string(hexcol.lstrip("#"))
        type_run = heading.add_run(f"[{r['scan_type']}] ")
        type_run.font.color.rgb = RGBColor.from_string("888888")
        heading.add_run(f"{idx}. {r['issue_title']}")

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        for key, val in info:
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = str(val)
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True

        if r["message"]:
            doc.add_paragraph(r["message"])

        p = doc.add_paragraph()
        p.add_run(flow_label).bold = True
        for fline in (r["data_flow"] or "(none)").split("\n"):
            cp = doc.add_paragraph()
            run = cp.add_run(fline)
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        if r["remediation"]:
            p = doc.add_paragraph()
            p.add_run("Remediation").bold = True
            _docx_add_markdown(doc, r["remediation"])

    if not rows:
        doc.add_paragraph("No issues found.")

    doc.save(path)


VALID_FORMATS = ("csv", "pdf", "docx")
FORMAT_ALIASES = {"all": VALID_FORMATS, "both": ("csv", "pdf")}


def parse_formats(fmt_arg):
    """Parse a comma-separated --format value into an ordered list of formats."""
    out = []
    for token in fmt_arg.lower().split(","):
        token = token.strip()
        if not token:
            continue
        for fmt in FORMAT_ALIASES.get(token, (token,)):
            if fmt not in VALID_FORMATS:
                raise SystemExit(
                    f"Unknown format '{fmt}'. Choose from: "
                    f"{', '.join(VALID_FORMATS)}, all, both (comma-separated)."
                )
            if fmt not in out:
                out.append(fmt)
    return out or ["csv"]


def resolve_outputs(output_arg, formats):
    """Map --output + parsed formats to concrete {format: path} targets."""
    base = output_arg or "snyk-results"
    base = re.sub(r"\.(csv|pdf|docx)$", "", base, flags=re.IGNORECASE)
    return {fmt: f"{base}.{fmt}" for fmt in formats}


def main():
    parser = argparse.ArgumentParser(
        description="Export Snyk Code and/or SCA findings (source-to-sink / dependency "
                    "path + remediation) to CSV, PDF and/or DOCX.",
        epilog="Pass extra Snyk CLI flags after '--', e.g. -- --org=acme --severity-threshold=high",
    )
    parser.add_argument("target", nargs="?", default=".",
                        help="Path to the project to scan (default: current dir)")
    parser.add_argument("-t", "--type", choices=("code", "sca", "both"), default="code",
                        help="Which Snyk tests to run: code (Snyk Code, default), "
                             "sca (open-source dependencies), or both")
    parser.add_argument("-f", "--format", default="csv",
                        help="Output format(s): csv, pdf, docx, all, or a comma-separated "
                             "list e.g. 'csv,docx' (default: csv)")
    parser.add_argument("-o", "--output",
                        help="Output path/base name (extension is set per format; "
                             "default: snyk-results)")
    parser.add_argument("--sarif-input",
                        help="Use an existing Snyk Code SARIF file instead of scanning "
                             "(implies --type code)")
    parser.add_argument("--project-root",
                        help="Root the SARIF paths are relative to "
                             "(default: target, or sarif file's dir with --sarif-input)")
    parser.add_argument("snyk_args", nargs="*",
                        help="Extra args passed through to the Snyk CLI")
    args = parser.parse_args()

    rows = []
    if args.sarif_input:
        if args.type != "code":
            print("Note: --sarif-input is a Snyk Code SARIF; ignoring --type "
                  f"{args.type} (SCA needs a live scan).", file=sys.stderr)
        sarif_path = validate_input_path(args.sarif_input)
        project_root = args.project_root or os.path.dirname(sarif_path) or "."
        with open(sarif_path, "r", encoding="utf-8") as fh:
            sarif = json.load(fh)
        rows += sarif_to_rows(sarif, SourceReader(project_root))
        scan_label = "Snyk Code (from SARIF)"
    else:
        project_root = args.project_root or args.target
        run_code = args.type in ("code", "both")
        run_sca = args.type in ("sca", "both")
        labels = []
        if run_code:
            sarif_path = run_snyk_code(args.target, args.snyk_args)
            with open(sarif_path, "r", encoding="utf-8") as fh:
                rows += sarif_to_rows(json.load(fh), SourceReader(project_root))
            labels.append("Snyk Code")
        if run_sca:
            sca_path = run_snyk_sca(args.target, args.snyk_args)
            if sca_path:
                with open(sca_path, "r", encoding="utf-8") as fh:
                    rows += sca_json_to_rows(json.load(fh))
                labels.append("Snyk Open Source (SCA)")
        scan_label = " + ".join(labels) or "Snyk"

    sort_rows(rows)
    targets = resolve_outputs(args.output, parse_formats(args.format))

    meta = {"project": os.path.abspath(project_root),
            "scan": scan_label,
            "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}
    writers = {"csv": write_csv, "pdf": write_pdf, "docx": write_docx}

    written = []
    for fmt, path in targets.items():
        writers[fmt](rows, path, meta)
        written.append(path)

    print(f"Wrote {len(rows)} issues ({summary_text(rows) or 'none'}) to "
          f"{', '.join(written)}", file=sys.stderr)


if __name__ == "__main__":
    main()
